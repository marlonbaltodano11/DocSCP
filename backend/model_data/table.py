from docx.table import Table as TableObject
from docx.table import _Row, _Cell
from docx.text.paragraph import Paragraph
from typing import Iterator, List, Tuple, Optional
from docx.document import Document
from docx.oxml import OxmlElement
from services.paragraph_service import ParagraphService
from docx.oxml.ns import qn
        
class Table:
    """Table data model."""
    
    def __init__(self, table: TableObject, document: Optional[Document] = None):
        self._table = table
        self._doc = document
    
    def _iter_row_cell(self, row: _Row) -> Iterator[Optional[_Cell]]:
        for _ in range(row.grid_cols_before):
            yield None
        for c in row.cells:
            yield c
        for _ in range(row.grid_cols_after):
            yield None
        
    def _iter_row_cell_texts(self, row: _Row) -> Iterator[str]:
        for cell in self._iter_row_cell(row):
            yield cell.text if cell else ""  # Devolver texto vacío si la celda es None

    def __str__(self) -> str:
        """Convert a table into a list of tuples containing cell texts."""
        return str([tuple(self._iter_row_cell_texts(row)) for row in self._table.rows])
    
    def _set_cell_background_color(self, cell, color_hex: str):
        """
        Sets the background color of a cell.

        :param cell: The cell to modify.
        :param color_hex: Background color in hex format (e.g., 'FFFF00' for yellow).
        """
        # Create a new XML element for the fill
        cell_element = cell._element
        cell_properties = cell_element.find('.//w:tcPr', cell_element.nsmap)
        
        if cell_properties is None:
            cell_properties = OxmlElement('w:tcPr')
            cell_element.append(cell_properties)

        # Add the fill color to the cell properties
        fill_element = OxmlElement('w:shd')
        fill_element.set(qn('w:fill'), color_hex)  # Set the fill color
        cell_properties.append(fill_element)
    
    def get_paragraphs(self) -> List[Paragraph]:
        """Extract paragraphs from the given table."""
        paragraphs: List[Paragraph] = []

        for row in self._table.rows:
            for cell in self._iter_row_cell(row):
                if cell:  # Si la celda no es None
                    paragraphs.extend(cell.paragraphs)

        return paragraphs

    def get_last_row(self) -> _Row:
        return self._table.rows[-1]
    
    def set_row(self, row: _Row, data: Optional[List[str]] = None, style: Optional[str] = None):        
        if data:
            # Limit the number of data items to the number of available cells
            num_cells = len(row.cells)
            for i, value in enumerate(data[:num_cells]):
                row.cells[i].text = ""
                default_paragraph = row.cells[i].paragraphs[0]
                ParagraphService.write_formatted_text(default_paragraph, value)
                
                # Apply style if provided
                if style and self._doc:
                    self._apply_style_to_cell(row.cells[i], style)
    
    def create_row(self, data: Optional[List[str]] = None, style: Optional[str] = None) -> _Row:
        """
        Creates a new row at the end of the table and fills it with the provided data.
        Optionally applies a style to the paragraphs in the cells.
        
        :param data: List of strings to populate the row's cells.
        :param style: Name of the style to apply to the text in the row's cells (optional).
        :return: The created row (_Row).
        """
        # Add a new row at the end of the table
        row = self._table.add_row()
        
        if data:
            # Limit the number of data items to the number of available cells
            num_cells = len(row.cells)
            for i, value in enumerate(data[:num_cells]):
                default_paragraph = row.cells[i].paragraphs[0]
                ParagraphService.write_formatted_text(default_paragraph, value)
                
                # Apply style if provided
                if style and self._doc:
                    self._apply_style_to_cell(row.cells[i], style)
        
        return row

    def create_subtitle_row(self, subtitle = "", style: Optional[str] = None, background_color: str = None) -> _Row:
        """
        Creates a new single cell row at the end of the table and fills it with the provided subtitle.
        Optionally applies a style to the paragraphs in the cells.
        
        :param subtitle: String to fill the row's cell.
        :param style: Name of the style to apply to the text in the row's cells (optional).
        :return: The created row (_Row).
        """
        # Add a new single cell row at the end of the table
        row = self._table.add_row()
        merged_cell = self.merge_cells_in_row(row)
        default_paragraph = merged_cell.paragraphs[0]
        ParagraphService.write_formatted_text(default_paragraph, subtitle)
        
        # Change background color if provided
        if background_color:
            self._set_cell_background_color(merged_cell, background_color)
        # Apply style if provided
        if style and self._doc:
            self._apply_style_to_cell(merged_cell, style)
        return row

    def _apply_style_to_cell(self, cell: _Cell, style: str) -> None:
        """
        Applies a style to all paragraphs within a cell.
        
        :param cell: Cell object where the style will be applied.
        :param style: Name of the style to apply.
        """
        ParagraphService.apply_style_to_paragraphs(cell.paragraphs, self._doc, style)
                
    def insert_row(self):
        """Insert a new row into a specific position of the table."""
        raise NotImplementedError("This method is not implemented yet.")
    
    def merge_cells_in_row(self, row: _Row, start_index: Optional[int] = None, end_index: Optional[int] = None) -> _Cell:
        """
        Combines a range of cells in a row into a single cell.
        
        If start_index and end_index are not specified, combines the entire row.

        :param row: The row containing the cells to merge.
        :param start_index: The index of the first cell to merge (0-based). Defaults to 0.
        :param end_index: The index of the last cell to merge (exclusive). Defaults to the length of the row.
        :raises IndexError: If start_index or end_index is out of range.
        :return: The merged cell.
        """
        row_length = len(row.cells)
        
        # Handle default values for indices
        start_index = start_index if start_index is not None else 0
        end_index = end_index if end_index is not None else row_length
        
        # Validate indices
        if start_index < 0 or start_index >= row_length:
            raise IndexError('start_index is out of range.')
        if end_index < 0 or end_index > row_length:
            raise IndexError('end_index is out of range.')
        if end_index <= start_index:
            raise ValueError('end_index must be greater than start_index.')

        # Merge cells in the specified range
        merged_cell = row.cells[start_index]
        for cell in row.cells[start_index + 1:end_index]:
            merged_cell = merged_cell.merge(cell)

        return merged_cell

        
