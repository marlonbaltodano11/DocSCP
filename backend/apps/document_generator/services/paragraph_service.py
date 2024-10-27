from docx.document import Document as DocumentObject
from docx.table import Table as TableObject
from docx.text.paragraph import Paragraph

class ParagraphService:
    """
    Clase que maneja la lógica de reemplazo de texto en párrafos y tablas.
    """
    def __init__(self, doc: DocumentObject):
        self.doc = doc

    def replace_text_in_paragraph(self, paragraph: Paragraph, search_text: str, replace_text: str):
        #if search_text in paragraph.text:
        for run in paragraph.runs:
            print(run.text)
            if search_text in run.text:
                run.text = run.text.replace(search_text, replace_text)

    def replace_text_in_table(self, table: TableObject, search_text: str, replace_text: str):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_text_in_paragraph(paragraph, search_text, replace_text)

    def replace_text(self, search_text: str, replace_text: str):
        # Reemplazar texto en párrafos
        for paragraph in self.doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, search_text, replace_text)

        # Reemplazar texto en tablas
        for table in self.doc.tables:
            self.replace_text_in_table(table, search_text, replace_text)